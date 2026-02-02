from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm

def export_to_docx(paper_data: dict, include_answers: bool = True) -> io.BytesIO:
    """
    Generates a Word document from the paper data.
    paper_data should contain: title, questions (list of dicts)
    """
    doc = Document()
    
    # Title
    title = doc.add_heading(paper_data.get("title", "Exam Paper"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle / Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Name: _______________  Score: _______")
    
    doc.add_paragraph("-" * 80)
    
    # Questions
    questions = paper_data.get("questions_snapshot", [])
    
    # Group by type for better layout? Or just sequential? 
    # Let's do sequential for now as they might be randomized intentionally.
    
    for idx, q in enumerate(questions, 1):
        # Question Stem
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"{idx}. [{q.get('q_type', 'Unknown')}] {q.get('content')}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        
        # Options (if choice)
        if q.get('q_type') in ['single', 'multi'] and q.get('options'):
            options = q.get('options')
            for opt in options:
                doc.add_paragraph(f"    {opt}", style='List Bullet')
        
        # Space for answer (if essay)
        if q.get('q_type') == 'essay':
            doc.add_paragraph("\n" * 5)
        
        # Space for others
        doc.add_paragraph()
        
    if include_answers:
        doc.add_page_break()
        doc.add_heading("Answer Key", level=1)
        for idx, q in enumerate(questions, 1):
            p_ans = doc.add_paragraph()
            p_ans.add_run(f"{idx}. {q.get('answer', 'N/A')}")
            if q.get("analysis"):
                p_ans.add_run(f"\n   Analysis: {q.get('analysis')}").italic = True

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_to_txt(paper_data: dict, include_answers: bool = True) -> io.BytesIO:
    title = paper_data.get("title", "Exam Paper")
    questions = paper_data.get("questions_snapshot", [])
    lines = [title, "", "Name: _______________  Score: _______", "", "-" * 60, ""]
    for idx, q in enumerate(questions, 1):
        lines.append(f"{idx}. [{q.get('q_type', 'Unknown')}] {q.get('content', '')}")
        if q.get("q_type") in ["single", "multi"] and q.get("options"):
            for opt in q.get("options") or []:
                lines.append(f"   - {opt}")
        lines.append("")
    if include_answers:
        lines.append("")
        lines.append("Answer Key")
        lines.append("-" * 60)
        for idx, q in enumerate(questions, 1):
            lines.append(f"{idx}. {q.get('answer', 'N/A')}")
            if q.get("analysis"):
                lines.append(f"   Analysis: {q.get('analysis')}")
    content = "\n".join(lines).encode("utf-8")
    stream = io.BytesIO(content)
    stream.seek(0)
    return stream

def export_to_pdf(paper_data: dict, include_answers: bool = True) -> io.BytesIO:
    title = paper_data.get("title", "Exam Paper")
    questions = paper_data.get("questions_snapshot", [])

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_x = 18 * mm
    margin_y = 18 * mm
    line_height = 5.5 * mm
    y = height - margin_y

    def new_page():
        nonlocal y
        c.showPage()
        y = height - margin_y

    def draw_line(text: str, font_name: str = "Helvetica", font_size: int = 11):
        nonlocal y
        if y < margin_y + line_height:
            new_page()
        c.setFont(font_name, font_size)
        c.drawString(margin_x, y, text)
        y -= line_height

    def wrap_text(text: str, max_chars: int = 80):
        text = text or ""
        chunks = []
        while len(text) > max_chars:
            cut = text.rfind(" ", 0, max_chars)
            if cut <= 0:
                cut = max_chars
            chunks.append(text[:cut].rstrip())
            text = text[cut:].lstrip()
        if text:
            chunks.append(text)
        return chunks or [""]

    draw_line(title, font_name="Helvetica-Bold", font_size=16)
    y -= line_height
    draw_line("Name: _______________  Score: _______", font_size=11)
    y -= line_height
    draw_line("-" * 90, font_size=10)
    y -= line_height

    for idx, q in enumerate(questions, 1):
        stem = f"{idx}. [{q.get('q_type', 'Unknown')}] {q.get('content', '')}"
        for line in wrap_text(stem, max_chars=90):
            draw_line(line, font_name="Helvetica-Bold", font_size=11)
        if q.get("q_type") in ["single", "multi"] and q.get("options"):
            for opt in q.get("options") or []:
                for line in wrap_text(f"   - {opt}", max_chars=90):
                    draw_line(line, font_size=10)
        y -= line_height

    if include_answers:
        new_page()
        draw_line("Answer Key", font_name="Helvetica-Bold", font_size=14)
        y -= line_height
        for idx, q in enumerate(questions, 1):
            ans = f"{idx}. {q.get('answer', 'N/A')}"
            for line in wrap_text(ans, max_chars=90):
                draw_line(line, font_size=11)
            if q.get("analysis"):
                for line in wrap_text(f"   Analysis: {q.get('analysis')}", max_chars=90):
                    draw_line(line, font_size=10)
            y -= line_height

    c.save()
    buffer.seek(0)
    return buffer
